import pandas as pd
import numpy as np
from matplotlib import pyplot as plt
import os
from datetime import datetime

from matplotlib import image as mpimg

from IPython.display import IFrame,clear_output

# for PDF reading
import textract

import re

import sys

import docx

from difflib import SequenceMatcher


#######################################################################################

def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()

#######################################################################################

def dms_to_dd(x,as_string=True):
    d,m,s = x.split()
    result = abs(float(d)) + float(m)/60. + float(s)/3600.
    if float(d) < 0:
        result = -result
    return result

#######################################################################################

def convert_state(state):
    return {'New Hampshire':'NH','Maine':'ME',
            'Massachusetts':'MA','New Hampshire/Maine':'NH'}[state]

#######################################################################################

def doy_to_date(x, year=2008, jan1=1):
    # jan1 is Day 1, usually
    #if np.isnan(x):
    #    return np.nan
    #print(x)
    result = ( pd.Period(year = year-1, month=12, day=31, freq='D') +  
                pd.to_timedelta(x+(1-jan1), unit='days') )
    return result.strftime('%Y-%m-%d')

#######################################################################################

def date_conversion(x, year=None, dateformat='%d-%m-%y'):

# year is Fall Year for date

# default interpretations:
# aaaa-bb-cc : Year/Month/Day

# PROBLEMATIC:
# aa-bb-cc : Month/Day/Year - or  Day/Month/Year if aa>12

# Returns string
    # Unknown / missing
    if np.any([True for i in ['(earliest/latest)', '-999','no data','no response',
                              'unknown', 'missing', 'unknown', 'unkown','none',
                              # the following are added for postcard data
                              #  2021-02-07
                              'died', 'no res','skip','omit','card not received',
                              'card not returned', 'moved','nursing home','delete'] 
                              if i in str(x).lower()]):
        return '-999'
    elif (str(x).strip()=='') | (str(x).strip()=='?') | (str(x).strip()=='-?-'):
        return '-999'
    elif x in ['0',0]:
        return '0'
    xx = str(x)
    
    if ('+1' in xx) | ('+2' in xx) | ('+3' in xx):
        xx = xx.split('+')[0].strip()
    
    outofbounds = False
    if ((year < 1678) | (year > 2262)) &  (year is not None):
        outofbounds = True
    
    if ((len(xx)==8) | ((len(xx)==10))) & ('-' not in xx) & ('/' not in xx):
        #print xx, year
        if (xx[-2]=='.') | ((len(xx)==8) & (xx.isdigit())):
            xx = '{}-{}-{}'.format(xx[:4],xx[4:6],xx[6:8]) # year, month, day
        #print xx, year
   
    try:
        if (len(xx)==8 ) & ('-' in xx):
            xdt = pd.to_datetime(xx, format=dateformat)
        else:
            xdt = pd.to_datetime(xx)
        d, m, y = xdt.day, xdt.month, xdt.year
    except ValueError as e:
        if (len(xx)==8) & ('-' in xx):
            # mostly a problem if 00-02-28 (i.e., thinking 00 is a month)
            if (xx[2]=='-') & (xx[5]=='-'):
                xx = '19'+xx
            else:
                xx = xx+', {}'.format(year)
        elif (len(xx)==10)& ('-' in xx) & outofbounds:
            if len(xx.split('-')[0]) >2:
                y,m, d = (int(i) for i in xx.split('-'))
            else:
                d,m,y = (int(i) for i in xx.split('-'))
            # latest thaw in August; earliest freeze in August
            if ((m<=8) & (y== year+1)) | ((m>=8) & (y==year)):
                return '{:04d}-{:02d}-{:02d}'.format(y,m,d)
            else:
                print ('+++++PROBLEM+++++')
                print(xx)
                xx = xx+', {}'.format(year)
            
        else:
            xx = xx+', {}'.format(year)
        try:
            xdt = pd.to_datetime(xx)
            d, m, y = xdt.day, xdt.month, xdt.year
        except ValueError as e:
            print ('**************')
            print (e)
            print ('    {} can not be converted to YYYY/MM/DD'.format(str(x)))
            print ('**************\n')
            return '-999'
            
    if year is not None:
        # print type(y), type(year)
        # latest thaw in September!, 
        # latest thaw in August; earliest freeze in August
        if ((m < 8) & (y != (year+1))) | ((m>9) & (y!=year)) | (
                ((m==8) | (m==9)) & (y!=year) & (y!=(year+1) ) ):
            if m<=8:
                yearnew = year+1
            else:
                yearnew = year+0
            
            print ('==================')
            print ('Wrong Year in table')
            print ('\tData from table: {} (start_year is {})'.format(xx, year))
            print ('\t\tYMD: {}-{:02d}-{:02d}'.format(y,m,d))
            print ('   Recorded (or added) ice date year {} should be {}\n'.format(y, yearnew))
            if (np.abs(int(y) - int(yearnew)) % 100) == 0:
                print ('\tFORCING YEAR TO NEW VALUE (wrong century)')
                y = yearnew
            # OTHERWISE TRY FIXING IT BY INVERTING DATE
            elif (len(xx)==8) & ('-' in xx):
                #print xx
                xx = '-'.join(xx.split('-')[::-1])
                #print xx
                # assuming default as before but switching backwards
                xdt = pd.to_datetime(xx,format=dateformat)
                d, m, y = xdt.day, xdt.month, xdt.year
                if ((m <= 8) & (y != year+1)) | ((m>8) & (y!=year)):
                    if m<=8:
                        yearnew = year+1
                    else:
                        yearnew = year
                    if (np.abs(int(y) - int(yearnew)) % 100) == 0:
                        print ('\tFORCING YEAR TO NEW VALUE (wrong century)')
                        y = yearnew
                    else:
                        print (x, xx)
                        print ('\tSTILL A PROBLEM. Recorded year {} should be {}'.format(y, yearnew))
                else:
                    print ('Problem fixed')
            else:
                print ('\tFORCING ICE YEAR TO NEW VALUE (assuming typo)')
                y = yearnew
            print ('   {}-{}, new corrected ice date {:}-{:02d}-{:02d}'.format(year, year+1,y,m,d))
    
    try:
        ##return '{:02d}-{:02d}-{:04d}'.format(m,d,y)
        return '{:04d}-{:02d}-{:02d}'.format(y,m,d)
    except ValueError as e:
        print ('*****FINAL*****')
        print (e)
        print ('**************')
        print ('{} can not be converted to YYYY/MM/DD'.format(str(x)))
        return '-999'



#######################################################################################
######## READ IN FILES ################################################################
#######################################################################################


def read_all_files(filename_dict, readin_dict , verbose=False,logfile=None, record_contributor=True):
    """
 INPUT: filename_dict is dictionary of files names, sorted by file type
        readin_dict is a list of corrections and column renames, etc. by filename

 OUTPUT: All files merged into a Pandas DataFrame 
    """

    default_ext = {
            'txt':{'delimiter':'\t'},
            'tab':{'delimiter':'\t'}
            }
    dfresult = pd.DataFrame()
# run through the files 
    for file_ext in filename_dict.keys():
        for f in filename_dict[file_ext]:
            
        
            default_values = {'header':0, 'delimiter':None, 'sheetname':False,
                      'lakename':None, 'city':None, 'state':None,'contributor':None, 'reorient':False,
                     'column_rename':None,'ncolumns':None, 'split':False,
                     'multi':False, 'index_col':None}
                     
            if file_ext in default_ext:
                for key, value in default_ext[file_ext].items():
                    default_values[key] = value
            if logfile is not None:
                logfile.write('===========\nReading in {}\n'.format(f))
            if (np.array([i in f for i in readin_dict.keys()])).any():
                lakeid = [i for i in readin_dict.keys() if i in f]
                if len(lakeid) > 1:
                    print ('WARNING. There are too many similarly named readin_dict items. Could be a problem.')
                    if logfile is not None:
                        logfile.write('\nWARNING. There are too many similarly named readin_dict items.\n')
                    break
                foo = readin_dict[lakeid[0]]
                for key,value in foo.items():
                    default_values[key] = value
            #if 'Updated Data 2019.5' in f:
            #    print(f)
            df = read_ts(f,delimiter=default_values['delimiter'], 
                   sheetname=default_values['sheetname'],
                   header=default_values['header'],
                   ncolumns=default_values['ncolumns'],
                   index_col=default_values['index_col'],
                   logfile = logfile,record_contributor=record_contributor)

            if verbose:
                if len(df)>0:
                    sys.stdout.write('\r[ {:150s} ]\r'.format(f))
                    #sys.stdout.flush()
                else:
                    sys.stdout.write('Skipping {}\n'.format(f))
                    #sys.stdout.flush()
                    
            # specific case for Maine lakes
            if default_values['reorient']:
                if logfile is not None:
                    logfile.write('\tReorienting table.\n')
                contributor = df.Contributor.values[0]
                #df = df.set_index(df.columns[0])
                #print('Maine drop')
                #display(df.head())
                #print(df.columns)
                df = df.drop('Contributor',axis=1,level=0).unstack().reset_index()
                #print('END Maine drop')
                df['Contributor'] = contributor

            if default_values['column_rename'] is not None:
                if logfile is not None:
                    logfile.write('\tRenaming columns.\n')
                df = df.rename(default_values['column_rename'],axis=1)
        
            if default_values['lakename'] is not None:
                if logfile is not None:
                    logfile.write('\tSetting lakename to {}\n'.format(default_values['lakename']))
                df['lake'] = default_values['lakename']
                
            if default_values['city'] is not None:
                if logfile is not None:
                    logfile.write('\tSetting city to {}\n'.format(default_values['city']))
                df['city'] = default_values['city']
                
            if default_values['state'] is not None:
                if logfile is not None:
                    logfile.write('\tSetting state to {}\n'.format(default_values['state']))
                df['state'] = default_values['state']
            
            if default_values['split']:
                # rearrange years/seasons
                if logfile is not None:
                    logfile.write('\tRearranging years/seasons\n')
                df = sort_by_season(df)
                
            if default_values['multi']:
                if logfile is not None:
                    logfile.write('\tSorting by events.\n')
                df = sort_by_events(df)
                
            #if default_values['lakename'] is not None:
            #    df['lake'] = default_values['lakename']
        
            if default_values['contributor'] is not None:
                if logfile is not None:
                    logfile.write('\tAssigning contributor: {}\n'.format(default_values['contributor']))
                df['Contributor'] = default_values['contributor']
        
            
            if 'Updated Data' in f:
                updated_year = f.split('Updated Data')[1].split('/')[0].strip()
                if updated_year == '2018':
                    updated_year = 2018.5
                elif updated_year == '':
                    updated_year = 2018.0
                else:
                    updated_year = float(updated_year)
                df['Updated Year'] = updated_year
                """
                if 'Updated Data 2020.5' in f:
                    df['Updated Year'] = 2020.5
                elif 'Updated Data 2020' in f:
                    df['Updated Year'] = 2020.0
                elif 'Updated Data 2019.5' in f:
                    df['Updated Year'] = 2019.5
                elif 'Updated Data 2018' in f:
                    df['Updated Year'] = 2018.5
                elif 'Updated Data 2019' in f:
                    df['Updated Year'] = 2019.0
                elif 'Updated Data' in f:
                    df['Updated Year'] = 2018.0
                """
                df['FileName'] = f
            try:
                dfresult = dfresult.append(df,ignore_index=True, sort=False)
            except:
                display(df)
                print(kasdf)
            
    return dfresult

#######################################################################################


def sort_by_events(df):
# Move multi-freeze thaw years into separate rows
    iceon1col = [c for c in ['Freeze date 1',] if c in df.columns][0]
    iceon2col = [c for c in ['Freeze date 2',] if c in df.columns][0]
    iceoff1col = [c for c in ['Thaw date 1',] if c in df.columns][0]
    iceoff2col = [c for c in ['Thaw date 2',] if c in df.columns][0]
    ind = ((~df[iceon1col].isnull() | ~df[iceoff1col].isnull()) &
            (~df[iceon2col].isnull() | ~df[iceoff2col].isnull()))
    # .copy
    dfoo = df[ind].copy()
    dfoo[iceon1col] = dfoo[iceon2col]
    dfoo[iceoff1col] = dfoo[iceoff2col]
    
    #print('sort by events Drop')
    df = df.append(dfoo,ignore_index=True,sort=False).drop([iceoff2col,iceon2col],axis=1)
    #print('END sort by events Drop')
    # display(df)
    return df


#######################################################################################


def sort_by_season(df):
    #print (df.columns)
    #display(df)
    yearcolumn = [c for c in ['Year','year'] if c in df.columns][0]
    iceoncolumn = [c for c in ['datefirstice','IceOnDOY','Ice On','Ice-On','Ice on'] if c in df.columns][0]
    iceoffcolumn = [c for c in ['datelastice','IceOffDOY','Ice Off','Ice-Off','Ice off'] if c in df.columns][0]

    # print df.columns
    lakecolumn = [c for c in ['lakeid','lake'] if c in df.columns][0]
    
    dropcolumns = [iceoncolumn, iceoffcolumn]
    
    dfresult = pd.DataFrame()
    for name, group in df.groupby(lakecolumn):
        iceoff = group[iceoffcolumn].tolist() + [np.nan]
        iceon = [np.nan] + group[iceoncolumn].tolist()
        try:
            years = [float(group[yearcolumn].astype(str).min()) - 1] + group[yearcolumn].tolist()
        except:
            print(yearcolumn)
            display(group[yearcolumn])
            display(df)
            #print (kmtpasdf)
        dfoo = pd.DataFrame({lakecolumn:name,
                            'Fall Year': years, 
                            iceoncolumn:iceon, 
                            iceoffcolumn:iceoff})
        dfresult = dfresult.append(dfoo, ignore_index=True,sort=False)
    #print('sort by season Drop')
    dfresult = dfresult.merge(df.drop(dropcolumns,axis=1), left_on=[lakecolumn,'Fall Year'], 
        right_on=[lakecolumn,yearcolumn], how='left')
    #print('END sort by season Drop')
    for c in dfresult.columns:
        ## if c not in [lakecolumn, yearcolumn,'Fall Year']+dropcolumns:
        if c in ['Contributor','Clerk']:
            ## print 'backfilling', c
            dfresult[c] = dfresult[c].fillna(method='bfill')
    
    
    ## clean up, remove no result years OK
    # print dfresult.shape
    ind = dfresult[iceoncolumn].isnull() & dfresult[iceoffcolumn].isnull()
    ## display(dfresult[ind])
    #.copy
    dfresult = dfresult[~ind].copy()
    #print dfresult.shape
    # remove duplicates
    #display(dfresult[dfresult.duplicated(subset=[lakecolumn,yearcolumn,
    #                                            iceoncolumn,iceoffcolumn],keep=False)])
    dfresult = dfresult.drop_duplicates(subset=[lakecolumn,yearcolumn,
                                                iceoncolumn,iceoffcolumn])    
    #print dfresult.shape
    if 'Duration' in dfresult.columns:
        #display(dfresult.tail(6))
        #display(df.tail(6))
        dfresult.loc[dfresult.index[:-1],'Duration'] = df.loc[df.index[:],'Duration'].values
        # last duration should be removed
        dfresult.loc[dfresult.index[-1],'Duration'] = np.nan
        
        if dfresult.lake.values[0]!='Mirror Lake':
            print(dfresult.columns)
            display(dfresult.head())
            print(brokend)
            
    return dfresult

#######################################################################################
#######################################################################################
#######################################################################################



def read_ts(filename, header=0, sheetname=False, index_col=None, logfile=None,delimiter=None,ncolumns=None,
    record_contributor=True):

    """ ncolumns : number of columns to keep, starting with first
    """

    filetype = filename.split('.')[-1].lower()
    if filetype == 'pdf':
        tsdf = read_pdf(filename,logfile=logfile)
    #elif filetype == 'jpg':
    #    tsdf = read_jpg(filename)
    elif filetype in ['csv','txt','tab']:
        tsdf = read_csv(filename, delimiter=delimiter, header=header,record_contributor=record_contributor)
    #elif filetype in ['txt']:
    #    tsdf = read_csv(filename, delimiter=delimiter, header=header)
    elif filetype in ['xls','xlsx']:
        tsdf = read_excel(filename, sheetname=sheetname, logfile=logfile, index_col=index_col,header=header,ncolumns=ncolumns,
                    record_contributor=record_contributor)
    elif filetype in ['doc','docx']:
        if 'Updated Data 2019.5' in filename:
            doc = docx.Document(filename)
            if logfile is not None:
                for p in doc.paragraphs:
                    logfile.write('\t{}\n'.format(p.text))
        tsdf = pd.DataFrame()
        """
        if 'Updated Data 2019.5' in filename:
            doc = docx.Document(filename)
            print ('=====================')
            print (filename)
            print ('=====================')
            for p in doc.paragraphs:
                print (p.text)
        """
    elif filetype in ['jpg']:
        if logfile is not None:
            logfile.write('\tSKIPPING\n')
        tsdf = pd.DataFrame()
    else:
        if logfile is not None:
            logfile.write('\tSKIPPING\n')
        tsdf = pd.DataFrame()
    return tsdf

#######################################################################################

def read_csv(filename, delimiter=None, encoding='utf-8', header=0, record_contributor=True):
    try:
        df = pd.read_csv(filename, delimiter=delimiter, encoding='utf-8',engine='python',header=header)
        if df.shape[1]==1:
            print('{}\n\tToo few columns. Trying a different method.'.format(filename))
            df = pd.read_csv(filename, delimiter=delimiter, encoding='utf-8',engine='c',header=header)
            print('New shape:',df.shape)
    except UnicodeDecodeError as e:
        df = pd.read_csv(filename, delimiter=delimiter, encoding='latin1',engine='python',header=header)
    contributor = filename.split('/')[-2]
    # remove comment line if it exists
    
    if df.iloc[0,0] == '#':
        #print('CSV # Drop')
        df = df.drop(0,axis=0)
        #print('END csv # Drop')
    if record_contributor:
        df['Contributor'] = contributor
    return df

#######################################################################################

def read_jpg(filename):
    text2 = textract.process(filename, encoding='ascii',  
            method='tesseract',layout=True).decode("utf8")

#######################################################################################

def read_excel(filename, header=0, sheetname=False, index_col=None, logfile=None, ncolumns=None,
        record_contributor = True):

    df = pd.read_excel(filename, header=header, sheet_name= sheetname,
                index_col = index_col)
    contributor = filename.split('/')[-2]
    
    if ncolumns is not None:
        df = df.iloc[:,:ncolumns]
        
    # remove all blank columns
    df = df.dropna(how='all',axis=1)
    
    # remove row with '#'
    try:
        if len([True for i in df.iloc[0,:].tolist() if '#' in str(i)]) > 0:
            #print('excel # Drop')
            df = df.drop(0,axis=0)
            #print('END excel # Drop')
            if logfile is not None:
                logfile.write('\tDropping Row 0\n')
            #display(df.head(2))
    except:
        pass
    
    ## SPECIAL CASES
    if 'NHFRA' in filename:
        df.loc[0:3,:] = df.loc[0:3,:].ffill(axis=1)
        df.columns = df.loc[2,:]
        finalcolumns = df.iloc[4,0:4]
        #print('NHFRA Drop')
        df = df.drop([0,1,2,3,4],axis=0)
        #print('END NHFRA Drop')
        df2 = pd.DataFrame()
        for c in df.columns.unique():
            # .copy
            dfoo = df[c].copy()
            dfoo.columns = finalcolumns
            dfoo.loc[:, 'lake']= c
            df2 = df2.append(dfoo,ignore_index=True,sort=False).dropna()
        df = df2.reset_index(drop=True)
    elif 'Sapna_data' in filename:
        df = df.set_index('Winter').unstack()[::2].reset_index().merge(
            df.set_index('Winter').unstack()[1::2].reset_index(), left_index=True, right_index=True)
    elif 'ice_in_out' in filename:
        df = df.append(pd.read_excel(filename, header=header, 
                sheet_name='MN Ice In'),sort=False,ignore_index=True)  
    elif 'Serwy' in filename:
        dfoo = df.copy()
        dfoo.Winter = dfoo.Winter.replace('1986/1897','1896/1897')
        
        ind = dfoo['Ice cover_off (or ice cover on and off)'].astype(str).str.contains('until')
        dfoo2 = dfoo.loc[ind,:].copy()
        dfoo2['Ice cover_on'] = dfoo.loc[ind, 'Ice cover_off (or ice cover on and off)'].apply(lambda x: x.split('until')[0])
        dfoo2['Ice cover_off (or ice cover on and off)'] = dfoo.loc[ind, 'Ice cover_off (or ice cover on and off)'].apply(lambda x: x.split('until')[1].split('break')[0])
        ind = dfoo2['Ice cover_off (or ice cover on and off)'].astype(str).str.contains('10.04.1985')
        dfoo2.loc[ind, 'Ice cover_off (or ice cover on and off)'] = '1985-04-10'


        ind = dfoo['Ice cover_off (or ice cover on and off)'].astype(str).str.contains('break')
        dfoo3 = dfoo.loc[ind,:].copy()
        dfoo3['Ice cover_on'] = dfoo.loc[ind, 'Ice cover_off (or ice cover on and off)'].apply(lambda x: x.split('until')[1].split('break')[1])
        dfoo3['Ice cover_off (or ice cover on and off)'] = dfoo.loc[ind, 'Ice cover_off (or ice cover on and off)'].apply(lambda x: x.split('until')[2])

        ind = dfoo['Ice cover_on'].astype(str).str.contains('until')
        dfoo4 = dfoo.loc[ind,:].copy()
        dfoo4['Ice cover_on'] = dfoo.loc[ind, 'Ice cover_on'].apply(lambda x: x.split('until')[0])
        dfoo4['Ice cover_off (or ice cover on and off)']= dfoo.loc[ind,'Ice cover_on'].apply(lambda x: x.split('until')[1].split('break')[0].split('brek')[0])
        dfoo4[' ice cover'] = np.nan

        # Add missing data (ice on early but no ice off )
        ind = dfoo['Ice cover_on'].astype(str).str.contains('brek|break') & (dfoo['Ice cover_off (or ice cover on and off)']=='no data')
        dfoo5 = dfoo.loc[ind,:].copy()
        dfoo5['Ice cover_on'] = dfoo5['Ice cover_on'].apply(lambda x: (pd.to_datetime(x.split('until')[1].split('break')[0])+pd.to_timedelta('1 day')).strftime('%Y-%m-%d'))
        dfoo5[' ice cover'] = np.nan
        dfoo5['Ice cover_off (or ice cover on and off)'] = np.nan

        # Add missing data (ice off but no ice on early)
        ind = (dfoo['Ice cover_on'].astype(str).str.contains('no data') | dfoo['Ice cover_on'].isnull()) & (~dfoo['Ice cover_off (or ice cover on and off)'].isnull() & (dfoo['Ice cover_off (or ice cover on and off)']!='no data'))
        dfoo6 = dfoo.loc[ind,:].copy()
        dfoo6['Ice cover_off (or ice cover on and off)'] = dfoo6['Ice cover_off (or ice cover on and off)'].apply(lambda x: (pd.to_datetime(x.split('until')[0]) - pd.to_timedelta('1 day')).strftime('%Y-%m-%d'))
        dfoo6[' ice cover'] = np.nan
        dfoo6['Ice cover_on'] = np.nan

        ind = ~(dfoo['Ice cover_on'].astype(str).str.contains('until|break|brek') | dfoo['Ice cover_off (or ice cover on and off)'].astype(str).str.contains('until|break'))
        df = dfoo[ind].append(dfoo2).append(dfoo3).append(dfoo4).append(dfoo5).append(dfoo6).reset_index(drop=True)
    elif ('Postcard' in filename) & ('addres' in filename.lower()) & ('addresses for data' not in filename.lower()):
        # MN may have empty column that should be filled in 
        df.columns = df.columns.str.replace('Unnamed: 36','ICE-IN 2018')
        nopivot = [c for c in df.columns if 'ICE' not in c.upper()]
        df.columns = df.columns.str.replace('\'','19')
        dfoo = df.set_index(nopivot).stack().reset_index()
        column = [c for c in dfoo.columns if 'level' in str(c)][0]
        # year is fall year if ice-in ; otherwise it is spring year so need to subtract one
        dfoo['Fall Year'] = dfoo[column].apply(lambda x: int(''.join([s for s in str(x).replace('\'','19') if s.isdigit()])) if 'IN' in x else int(''.join([s for s in str(x).replace('\'','19') if s.isdigit()])) - 1)
        dfoo.loc[:, 'ICE-IN'] = dfoo.loc[dfoo[column].str.upper().str.contains('IN'), 0]
        dfoo.loc[:,'ICE-OUT'] = dfoo.loc[dfoo[column].str.upper().str.contains('OUT|OFF'), 0]
        df = dfoo.copy()
    elif ('Postcard' in filename) & ('fix_Doug' in filename):
        dfoo = df.loc[1:34,['Unnamed: 1']+df.columns.tolist()[3:-9:2]].rename({'Unnamed: 1':'Spring Year'},axis=1).set_index('Spring Year').unstack().reset_index().rename({0:'Ice Off','level_0':'lakename'},axis=1)
        dfoo2 = df.loc[45:78,['Unnamed: 1']+df.columns.tolist()[3:-9:2]].rename({'Unnamed: 1':'Fall Year'},axis=1).set_index('Fall Year').unstack().reset_index().rename({0:'Ice In','level_0':'lakename'},axis=1)
        df = dfoo.append(dfoo2, ignore_index=True)
        #display(dfoo.head())
        #display(dfoo2.head())

    df = df.dropna(how='all',axis=1)
    df = df.dropna(how='all',axis=0)

    if record_contributor:
        df['Contributor'] = contributor
    return df

#######################################################################################

def read_pdf(filename,logfile=None):
    try:
        text2 = textract.process(filename, encoding='ascii',  
            method='pdftotext',layout=True).decode("utf8")
    except:
        print('ERROR reading PDF: {}'.format(filename))
    # THESE ARE SPECIFIC CASES
    #  George and Hodkins New England lakes (ofr2010)
    if 'George' in filename:
        column1 = [t.strip('\x0c').split()[0] for t in text2.split('\n') if len(t.strip('\x0c').split())>0]
        column2 = [t.strip('\x0c').strip()[5:42].strip().split('   ')[0] for t in text2.split('\n') if len(t.strip('\x0c').split())>0]
        column4 = [t.strip('\x0c').strip()[-6:] for t in text2.split('\n') if len(t.strip('\x0c').split())>0]
        column3a = [t.strip('\x0c').strip()[38:83] for t in text2.split('\n') if len(t.strip('\x0c').split())>0]
        column3a = [[i.strip() for i in c.split('   ') if (len(i)>=7)] for c in column3a]
        column3 = []
        for c in column3a:
            if len(c)==0:
             column3.append('')
            else:
                column3.append(c[0])
        georgedata = {column1[0]:column1[1:], 
                        column2[0]:column2[1:], 
                        column3[0]:column3[1:], 
                        column4[0]:column4[1:],
                        'Lake':'George Lake',
                        'Contributor':'John Magnuson',
                        'Latitude':43.83,
                        'Longitude':-73.43}
        df = pd.DataFrame(georgedata)
        df['Year'] = df['Year'].astype(int)
        df['Fall Year'] = df['Year']-1
        # FIX READ-IN ERROR for missing data
        df.loc[df.Frozen.str.contains(','),'Frozen'] = np.nan
        df = df.sort_values('Fall Year')
    elif ('Winnipesaukee' in filename) | ('Sunapee' in filename):
        if 'Winnipesaukee' in filename:
            skipnrows = 9
            lakename = 'Winnipesaukee'
        else:
            skipnrows = 7
            lakename = 'Sunapee'
            
        df = pd.DataFrame(columns=['Spring Year','Ice In Date', 'Ice Out Date', 'Ice Cover Days'])
        for row in text2.replace('\x0c','').split('\n')[skipnrows:]:
            if len(row.split()) <= 6:
                years = row.split()[::2]
                iceouts= row.split()[1::2]
                if len(years)>len(iceouts):
                    iceouts.append(np.nan)
                df = df.append(pd.DataFrame({'Spring Year':years,'Ice Out Date':iceouts}),ignore_index=True,sort=False)
            elif len(row.split()) ==8:
                years = row.split()[:5:2]
                iceouts = row.split()[1:4:2] + row.split()[6:7]
                iceins = [np.nan,np.nan,row.split()[5]]
                durations = [np.nan,np.nan, row.split()[7]]
                df = df.append(pd.DataFrame({'Spring Year':years, 'Ice In Date':iceins, 'Ice Out Date':iceouts,'Ice Cover Days':durations}),ignore_index=True,sort=False)
        df = df.sort_values('Spring Year')
        df['Lake'] = lakename
        df['Contributor'] = 'Julia Daly'
    elif 'ofr2010' in filename:
        source_list = [re.split('Observers|Observer', text2.split('\x0cTable')[i])[1].split('\n')[0] for i in range(2,31)]
        comments_list = [' '.join(text2.split('\x0cTable')[i].split('Comments')[1].split('Year')[0].split())  
                 if (len(text2.split('\x0cTable')[i].split('Comments')) > 1) else '' for i in range(2,31)]
        longitude_list = [text2.split('\x0cTable')[i].split('Location of lake')[1][29:37] for i in range(2,31)]
        latitude_list = [text2.split('\x0cTable')[i].split('Location of lake')[1][9:17] for i in range(2,31)]
        latitude_list = [round(dms_to_dd('{} {} {}'.format(l[:2],l[2:4],l[5:7])),4) for l in latitude_list]
        longitude_list= [round(dms_to_dd('-{} {} {}'.format(l[:2],l[2:4],l[5:7])),4) for l in longitude_list]
        lakenames = [t.split('for')[1].split('.')[0].strip() for t in text2.split('\x0cTable')[2:]]
        tablenumbers = [int(float(t.split('Ice-out')[0])) for t in text2.split('\x0cTable')[2:]]
        tableslist = [t.split('Julian day')[-1:][0] for t in text2.split('\x0cTable')[2:]]
        year_list = [t.replace('\n',' ').replace('--','-999').strip().split()[:-1:2] for t in tableslist]
        doy_list = [t.replace('\n',' ').replace('--','-999').replace('NCIC','999').strip().split()[1:-1:2] for t in tableslist]
        year_list = [[int(d) for d in y] for y in year_list]
        doy_list = [[int(d) for d in y] for y in doy_list]

        df = pd.DataFrame()
        for lakename,tablenum,year, doy,latdd,londd,source,comments in zip(lakenames,tablenumbers, 
                                                          year_list, doy_list,
                                                          latitude_list, longitude_list,
                                                                   source_list,comments_list):
            nyears = len(year)
    #print len(year),len(doy)
    #print np.array(year).min(), np.array(year).max()
    #print 'Table {}. {} '.format(tablenum, lakename)
            dfoo = pd.DataFrame({'Fall Year': np.array(year)-1, 
                                'Year': year,
                                'Julian day':doy,
                                'Table':tablenum,
                                'Lake':lakename.split(',')[0],
                                'Latitude':latdd,
                                'Longitude':londd,
                                'Source':source,
                                'Comments':comments,
                                'State': convert_state(lakename.split(',')[1].strip()),
                                'Contributor':'Glenn Hodgkins'
                                })
            df = df.append(dfoo, ignore_index=True,sort=False)
        df = df.sort_values(['Table','Year'])
    else:
        if logfile is not None:
            logfile.write('\tSKIPPING. Not able to read PDF file\n')
        df = pd.DataFrame()
        
    return df
    


